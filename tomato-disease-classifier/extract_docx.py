import docx
import sys

def extract_text(filepath):
    doc = docx.Document(filepath)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fullText.append(cell.text)
                
    return '\n'.join(fullText)

if __name__ == '__main__':
    text = extract_text(sys.argv[1])
    with open('srs_content_raw.txt', 'w', encoding='utf-8') as f:
        f.write(text)
    print("Done")
